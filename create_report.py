"""
Script to generate the Practical 8 Word Document Report.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

# ═══════════════ TITLE PAGE ═══════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Practical 8', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Movie Review Sentiment Analysis Project', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
info_lines = [
    ('Name', 'Rishil Pawar'),
    ('Roll No.', 'C1-18'),
    ('Subject', 'Natural Language Processing (NLP) Lab'),
    ('Practical No.', '8'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = p.add_run(f'{label}: ')
    run_label.bold = True
    run_label.font.size = Pt(13)
    run_value = p.add_run(value)
    run_value.font.size = Pt(13)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GitHub Repository: ')
run.bold = True
run.font.size = Pt(12)
run2 = p.add_run('https://github.com/Rishil-Pawar/Movie-Review-Sentiment-Analysis')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0, 0, 255)

doc.add_page_break()

# ═══════════════ TABLE OF CONTENTS ═══════════════
add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Introduction / Aim',
    '2. Dataset Description',
    '3. Libraries & Dependencies',
    '4. Step 1: Install & Import Libraries',
    '5. Step 2: Load & Explore Dataset (EDA)',
    '6. Step 3: Text Preprocessing',
    '7. Step 4: Feature Extraction (TF-IDF)',
    '8. Step 5: Model Training & Evaluation',
    '9. Step 6: Comparative Analysis',
    '10. Step 7: Best Model Selection',
    '11. Step 8: Model Saving & Loading',
    '12. Step 9: GUI-based Prediction (Gradio)',
    '13. Conclusion',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ═══════════════ 1. INTRODUCTION ═══════════════
add_heading_styled('1. Introduction / Aim', level=1)
add_para(
    'The aim of this practical is to build a complete Movie Review Sentiment Analysis system '
    'using Natural Language Processing (NLP) techniques. The project uses the IMDB Dataset '
    'containing 50,000 movie reviews (25,000 positive and 25,000 negative) to train and '
    'evaluate multiple classification models. The best performing model is then integrated '
    'into a GUI-based application built with Gradio for real-time sentiment prediction.'
)
doc.add_paragraph()

# ═══════════════ 2. DATASET ═══════════════
add_heading_styled('2. Dataset Description', level=1)
add_para(
    'Dataset: IMDB Dataset (IMDB Dataset.csv)\n'
    '• Total Reviews: 50,000\n'
    '• Positive Reviews: 25,000\n'
    '• Negative Reviews: 25,000\n'
    '• Columns: review (text), sentiment (positive/negative)\n'
    '• Missing Values: 0\n'
    '• Source: Kaggle IMDB Dataset'
)
doc.add_paragraph()

# ═══════════════ 3. LIBRARIES ═══════════════
add_heading_styled('3. Libraries & Dependencies', level=1)
libs = [
    'pandas – Data manipulation and analysis',
    'numpy – Numerical computing',
    'nltk – Natural Language Toolkit (stopwords, stemming)',
    'scikit-learn – Machine Learning (TF-IDF, classifiers, metrics)',
    'matplotlib – Data visualization (charts)',
    'seaborn – Statistical data visualization (heatmaps)',
    'pickle – Model serialization',
    'gradio – GUI-based web interface for predictions',
]
for lib in libs:
    doc.add_paragraph(lib, style='List Bullet')
doc.add_paragraph()

# ═══════════════ 4. STEP 1 ═══════════════
add_heading_styled('4. Step 1: Install & Import Libraries', level=1)
add_para('All required libraries are installed using pip and imported:', bold=True)
code = '''!pip install nltk gradio scikit-learn pandas matplotlib seaborn -q

import pandas as pd
import numpy as np
import nltk, re, warnings, pickle
import matplotlib.pyplot as plt
import seaborn as sns

warnings.filterwarnings('ignore')
nltk.download('stopwords')
nltk.download('punkt')

from nltk.corpus import stopwords
from nltk.stem import PorterStemmer

print("All libraries imported successfully.")'''
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(9)
add_para('Output: All libraries imported successfully.', italic=True)
doc.add_paragraph()

# ═══════════════ 5. STEP 2 ═══════════════
add_heading_styled('5. Step 2: Load & Explore Dataset (EDA)', level=1)
add_para('The IMDB dataset is loaded and explored:', bold=True)
code2 = '''df = pd.read_csv('/content/IMDB Dataset.csv')
print("Dataset Shape:", df.shape)
print("Class Distribution:")
print(df['sentiment'].value_counts())
print("Missing Values:", df.isnull().sum().sum())'''
p = doc.add_paragraph()
run = p.add_run(code2)
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_para('\nOutput:', italic=True)
add_para('Dataset Shape: (50000, 2)')
add_para('Class Distribution: positive 25000, negative 25000')
add_para('Missing Values: 0')

add_para('\nVisualization: Class Distribution (Bar Chart & Pie Chart)', bold=True)
add_para('A bar chart and pie chart are generated showing the balanced distribution of positive and negative reviews in the dataset. (See Colab notebook for output screenshot)', italic=True)
doc.add_paragraph()

# ═══════════════ 6. STEP 3 ═══════════════
add_heading_styled('6. Step 3: Text Preprocessing', level=1)
add_para('Text preprocessing pipeline:', bold=True)
steps = [
    'Convert text to lowercase',
    'Remove HTML tags using regex',
    'Remove non-alphabetic characters',
    'Tokenize the text',
    'Remove stopwords (English)',
    'Apply Porter Stemmer for word stemming',
    'Filter out words with length ≤ 2',
    'Join tokens back into cleaned string',
    'Map sentiment labels: positive → 1, negative → 0',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

code3 = '''stemmer = PorterStemmer()
stop_words = set(stopwords.words('english'))

def preprocess(text):
    text = text.lower()
    text = re.sub(r'<.*?>', '', text)
    text = re.sub(r'[^a-z\\s]', '', text)
    tokens = text.split()
    tokens = [stemmer.stem(w) for w in tokens
              if w not in stop_words and len(w) > 2]
    return ' '.join(tokens)

df['clean_review'] = df['review'].apply(preprocess)
df['label'] = df['sentiment'].map({'positive': 1, 'negative': 0})'''
p = doc.add_paragraph()
run = p.add_run(code3)
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_para('\nOutput: Preprocessing complete! Total reviews processed: 50000', italic=True)
add_para('(Before/After preprocessing comparison shown in notebook)', italic=True)
doc.add_paragraph()

# ═══════════════ 7. STEP 4 ═══════════════
add_heading_styled('7. Step 4: Feature Extraction (TF-IDF)', level=1)
add_para('TF-IDF (Term Frequency – Inverse Document Frequency) vectorization:', bold=True)
code4 = '''from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split

X = df['clean_review']
y = df['label']

X_train_raw, X_test_raw, y_train, y_test = train_test_split(
    X, y, test_size=0.20, random_state=42, stratify=y)

tfidf = TfidfVectorizer(max_features=10000, ngram_range=(1, 2), min_df=2)
X_train = tfidf.fit_transform(X_train_raw)
X_test  = tfidf.transform(X_test_raw)'''
p = doc.add_paragraph()
run = p.add_run(code4)
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_para('\nOutput:', italic=True)
add_para('Training set size: (40000, 10000)')
add_para('Test set size: (10000, 10000)')
add_para('Vocabulary size: 10000')
doc.add_paragraph()

# ═══════════════ 8. STEP 5 ═══════════════
add_heading_styled('8. Step 5: Model Training & Evaluation', level=1)
add_para('Five classification models are trained and evaluated:', bold=True)

models_info = [
    ('Logistic Regression', 'LogisticRegression(max_iter=1000, C=1.0)'),
    ('Multinomial Naive Bayes', 'MultinomialNB(alpha=0.1)'),
    ('Linear SVM', 'LinearSVC(max_iter=2000, C=1.0)'),
    ('Random Forest', 'RandomForestClassifier(n_estimators=100, random_state=42)'),
    ('Gradient Boosting', 'GradientBoostingClassifier(n_estimators=100, random_state=42)'),
]
for name, config in models_info:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(config)

add_para('\nEach model is trained on the TF-IDF features and evaluated using Accuracy, '
         'Precision, Recall, and F1-Score metrics on the test set.')
doc.add_paragraph()

# ═══════════════ 9. STEP 6 ═══════════════
add_heading_styled('9. Step 6: Comparative Analysis', level=1)
add_para('Model Comparison Table:', bold=True)

# Create comparison table
table = doc.add_table(rows=6, cols=6)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Model', 'Accuracy %', 'Precision %', 'Recall %', 'F1-Score %', 'Time (s)']
data = [
    ['Logistic Regression', '89.05', '88.43', '89.86', '89.14', '1.27'],
    ['Linear SVM', '88.77', '88.17', '89.56', '88.86', '3.36'],
    ['Multinomial Naive Bayes', '86.35', '85.00', '88.28', '86.61', '0.08'],
    ['Random Forest', '85.00', '85.21', '84.70', '84.95', '126.87'],
    ['Gradient Boosting', '81.27', '78.25', '86.62', '82.22', '164.44'],
]
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for row_idx, row_data in enumerate(data, 1):
    for col_idx, value in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = value

doc.add_paragraph()
add_para('Key Observations:', bold=True)
observations = [
    'Logistic Regression achieves the highest accuracy (89.05%) and best F1-Score (89.14%).',
    'Linear SVM performs comparably with 88.77% accuracy.',
    'Multinomial Naive Bayes is the fastest model (0.08s) with 86.35% accuracy.',
    'Random Forest and Gradient Boosting are significantly slower without accuracy gains.',
    'All models show balanced precision and recall values.',
]
for obs in observations:
    doc.add_paragraph(obs, style='List Bullet')

add_para('\nVisualization:', bold=True)
add_para('• Horizontal bar chart comparing model accuracies')
add_para('• Grouped bar chart comparing Precision, Recall, and F1-Score')
add_para('• Confusion matrices (heatmaps) for top 3 models')
add_para('(See Colab notebook for output screenshots)', italic=True)
doc.add_paragraph()

# ═══════════════ 10. STEP 7 ═══════════════
add_heading_styled('10. Step 7: Best Model Selection', level=1)
add_para('Based on the comparative analysis, Logistic Regression is selected as the best model:', bold=True)
doc.add_paragraph()

best_table = doc.add_table(rows=5, cols=2)
best_table.style = 'Light Grid Accent 1'
best_data = [
    ('Model', 'Logistic Regression'),
    ('Accuracy', '89.05%'),
    ('Precision', '88.43%'),
    ('Recall', '89.86%'),
    ('F1-Score', '89.14%'),
]
for idx, (label, value) in enumerate(best_data):
    best_table.rows[idx].cells[0].text = label
    best_table.rows[idx].cells[1].text = value
    for paragraph in best_table.rows[idx].cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph()
add_para('Detailed Classification Report:', bold=True)
report = '''              precision    recall  f1-score   support

    Negative       0.90      0.88      0.89      5000
    Positive       0.88      0.90      0.89      5000

    accuracy                           0.89     10000
   macro avg       0.89      0.89      0.89     10000
weighted avg       0.89      0.89      0.89     10000'''
p = doc.add_paragraph()
run = p.add_run(report)
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_para('\nNote: While the accuracy is 89.05% (below the 95% threshold mentioned in guidelines), '
         'Logistic Regression is the best performing model among all trained classifiers. '
         'To achieve ≥95% accuracy, advanced techniques such as deep learning (LSTM, BERT) '
         'or more sophisticated feature engineering would be required.', italic=True)
doc.add_paragraph()

# ═══════════════ 11. STEP 8 ═══════════════
add_heading_styled('11. Step 8: Model Saving & Loading', level=1)
add_para('The best model and TF-IDF vectorizer are saved using Python pickle:', bold=True)
code5 = '''import pickle

with open('best_model.pkl', 'wb') as f:
    pickle.dump(best_model, f)

with open('tfidf_vectorizer.pkl', 'wb') as f:
    pickle.dump(tfidf, f)

# Verification
with open('best_model.pkl', 'rb') as f:
    loaded_model = pickle.load(f)
with open('tfidf_vectorizer.pkl', 'rb') as f:
    loaded_tfidf = pickle.load(f)

test_pred = loaded_model.predict(
    loaded_tfidf.transform(["This movie was great!"]))
print(f"Reload test — 'This movie was great!' → "
      f"{'Positive' if test_pred[0]==1 else 'Negative'}")'''
p = doc.add_paragraph()
run = p.add_run(code5)
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_para("\nOutput: Reload test — 'This movie was great!' → Positive", italic=True)
add_para("Models loaded and verified successfully.", italic=True)
doc.add_paragraph()

# ═══════════════ 12. STEP 9 ═══════════════
add_heading_styled('12. Step 9: GUI-based Prediction (Gradio)', level=1)
add_para('A Gradio-based web interface is built for real-time sentiment prediction:', bold=True)
code6 = '''import gradio as gr

def predict_sentiment(review: str) -> str:
    if not review.strip():
        return "Please enter a review."
    cleaned = preprocess(review)
    vec = tfidf.transform([cleaned])
    pred = best_model.predict(vec)[0]
    try:
        prob = best_model.predict_proba(vec)[0]
        conf = round(max(prob) * 100, 1)
        conf_str = f"  (confidence: {conf}%)"
    except AttributeError:
        conf_str = ""
    label = "POSITIVE" if pred == 1 else "NEGATIVE"
    return f"{label}{conf_str}"

demo = gr.Interface(
    fn=predict_sentiment,
    inputs=gr.Textbox(lines=5, placeholder="Paste your movie review here...",
                      label="Movie Review"),
    outputs=gr.Textbox(label="Predicted Sentiment"),
    title="Movie Review Sentiment Analyser",
    description="Classifies movie reviews as Positive or Negative.",
    examples=[
        ["This movie was absolutely brilliant!"],
        ["Terrible film. Complete waste of time."],
    ],
    theme=gr.themes.Soft(),
    allow_flagging="never"
)
demo.launch(share=True)'''
p = doc.add_paragraph()
run = p.add_run(code6)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
add_para('GUI Features:', bold=True)
features = [
    'Text input box for entering movie reviews',
    'Predicted sentiment output with confidence percentage',
    'Pre-loaded example reviews for quick testing',
    'Clean, modern Soft theme interface',
    'Share link for external access',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

add_para('\n(See Colab notebook for Gradio GUI output screenshot)', italic=True)
doc.add_paragraph()

# ═══════════════ 13. CONCLUSION ═══════════════
add_heading_styled('13. Conclusion', level=1)
add_para(
    'In this practical, a complete Movie Review Sentiment Analysis system was built using '
    'NLP techniques. The following key steps were performed:'
)
conclusions = [
    'Dataset Loading & EDA: The IMDB dataset (50,000 reviews) was loaded and explored, confirming balanced class distribution.',
    'Text Preprocessing: Reviews were cleaned using lowercasing, HTML tag removal, stopword removal, and Porter Stemming.',
    'Feature Extraction: TF-IDF vectorization with 10,000 features and bigrams was used to convert text to numerical features.',
    'Model Training: Five models (Logistic Regression, Naive Bayes, Linear SVM, Random Forest, Gradient Boosting) were trained and evaluated.',
    'Comparative Analysis: All models were compared using Accuracy, Precision, Recall, and F1-Score metrics.',
    'Best Model: Logistic Regression was selected as the best model with 89.05% accuracy and 89.14% F1-Score.',
    'Model Persistence: The best model and vectorizer were saved using pickle for reuse.',
    'GUI Implementation: A Gradio-based web interface was built for real-time sentiment prediction with confidence scores.',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Number')

doc.add_paragraph()
add_para('GitHub Repository Link: https://github.com/Rishil-Pawar/Movie-Review-Sentiment-Analysis', bold=True)

# ═══════════════ SAVE ═══════════════
doc.save(r'c:\Users\Rishil\Documents\Movie Reviews\Practical_8_Movie_Review_Sentiment_Analysis_Report.docx')
print("Report saved successfully!")
